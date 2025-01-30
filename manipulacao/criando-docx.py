from docx import Document
#criação de documento
doc = Document()

doc.add_heading('Titulo Principal', level= 1)
doc.add_paragraph('Este é um paragráfo de exemplo')
doc.save('exemplo.docx')

#leitura do documento

doc = Document('exemplo.docx')
for paragraph in doc.paragraphs: #percorre o documento como se fosse uma lista 
    print(paragraph.text)

#adicionar tabelas ao documento

doc = Document()
table = doc.add_table(rows=2, cols=2)
table.cell(0,0).text = 'cabeçalho 1'
table.cell(0,1).text = 'cabeçalho 2'
table.cell(1,0).text = 'valor 1'
table.cell(1,1).text = 'valor 2'

doc.save('table.docx')